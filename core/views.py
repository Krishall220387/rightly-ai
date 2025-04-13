from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Blog, Document
from django.http import JsonResponse, StreamingHttpResponse, HttpResponse, HttpResponseForbidden
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
import json
import time
from .utils import extract_text_from_file, generate_blog_content, check_grammar
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from django.contrib.auth import login as auth_login
import openai
import os
from django.conf import settings
from .forms import UserRegistrationForm, DocumentUploadForm, BlogCreationForm
from django.urls import reverse
from openai import OpenAI
import logging
import PyPDF2
import docx2txt
from rest_framework.decorators import api_view

logger = logging.getLogger(__name__)

def home(request):
    if request.user.is_authenticated:
        return redirect('core:dashboard')
    return render(request, 'core/landing.html')

def register(request):
    """
    Handle user registration with form validation and success/error messages.
    """
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            messages.success(request, 'Registration successful! Welcome to Rightly.ai!')
            return redirect('core:dashboard')
        else:
            for error in form.errors.values():
                messages.error(request, error)
    else:
        form = UserRegistrationForm()
    return render(request, 'core/register.html', {'form': form})

def user_login(request):
    """
    Handle user login with validation and messages.
    """
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, f'Welcome back, {username}!')
            return redirect('core:dashboard')
        else:
            messages.error(request, 'Invalid username or password.')
    return render(request, 'core/login.html')

@login_required
def user_logout(request):
    """
    Handle user logout.
    """
    logout(request)
    messages.info(request, 'You have been logged out successfully.')
    return redirect('core:login')

@login_required
def dashboard(request):
    """
    Display user dashboard with documents and blogs.
    """
    documents = Document.objects.filter(user=request.user).order_by('-uploaded_at')
    blogs = Blog.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'core/dashboard.html', {
        'documents': documents,
        'blogs': blogs
    })

@login_required
def upload_document(request):
    """
    Handle document upload with form validation and file processing.
    """
    if request.method == 'POST':
        try:
            # Get the uploaded file and title
            uploaded_file = request.FILES.get('file')
            title = request.POST.get('title', uploaded_file.name)
            
            if not uploaded_file:
                messages.error(request, 'No file was uploaded.')
                return redirect('core:upload_document')
            
            # Create document directory if it doesn't exist
            user_doc_dir = os.path.join('media', 'documents', f'user_{request.user.id}')
            os.makedirs(user_doc_dir, exist_ok=True)
            
            # Create and save the document
            document = Document(
                user=request.user,
                title=title,
                file=uploaded_file
            )
            document.save()
            
            messages.success(request, 'Document uploaded successfully!')
            return redirect('core:dashboard')
            
        except Exception as e:
            print(f"Error uploading document: {str(e)}")
            messages.error(request, f'Error uploading document: {str(e)}')
            return redirect('core:upload_document')
    
    return render(request, 'core/upload_document.html')

@login_required
def delete_blog(request, blog_id):
    """
    Handle blog deletion with success/error messages.
    """
    try:
        print(f"Attempting to delete blog with ID: {blog_id}")  # Debug log
        blog = get_object_or_404(Blog, id=blog_id, user=request.user)
        print(f"Found blog: {blog.blog_title}")  # Debug log
        blog.delete()
        print(f"Successfully deleted blog with ID: {blog_id}")  # Debug log
        return JsonResponse({'status': 'success', 'message': 'Blog deleted successfully!'})
    except Exception as e:
        print(f"Error deleting blog: {str(e)}")  # Debug log
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

@login_required
def create_blog(request):
    """Handle both GET and POST requests for blog creation."""
    if request.method == 'GET':
        # Get user's uploaded documents for selection
        documents = Document.objects.filter(user=request.user).order_by('-uploaded_at')
        return render(request, 'core/create_blog.html', {'documents': documents})
    
    # Handle POST request
    try:
        # Extract data from request
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            data = request.POST.dict()
            # Handle document_ids from form data
            document_ids = request.POST.getlist('selected_documents', [])
            data['document_ids'] = json.dumps(document_ids)

        topic = data.get('topic')
        tone = data.get('tone', 'professional')
        user_keywords = [kw.strip() for kw in data.get('user_keywords', '').split(',') if kw.strip()]
        document_ids = json.loads(data.get('document_ids', '[]'))

        # Validate required fields
        if not topic:
            return JsonResponse({'error': 'Topic is required'}, status=400)
        
        if not document_ids:
            return JsonResponse({'error': 'Please select at least one document'}, status=400)

        # Process selected documents
        documents_content = []
        successful_extractions = 0
        
        try:
            documents = Document.objects.filter(id__in=document_ids, user=request.user)
            if not documents.exists():
                return JsonResponse({'error': 'No valid documents selected'}, status=400)
            
            for doc in documents:
                try:
                    print(f"Processing document: {doc.title}")
                    content = extract_text_from_file(doc.file)
                    
                    if content and content not in ["Error extracting text from file", 
                                                 "No content could be extracted",
                                                 "No text could be extracted from PDF",
                                                 "No text could be extracted from DOCX"]:
                        documents_content.append(f"\n\nDocument: {doc.title}\n{content}")
                        successful_extractions += 1
                    else:
                        print(f"No content extracted from document: {doc.title}")
                        
                except Exception as e:
                    print(f"Error processing document {doc.title}: {str(e)}")
                    continue
            
            if successful_extractions == 0:
                return JsonResponse({
                    'error': 'Could not extract content from any of the selected documents. Please ensure the documents contain readable text.'
                }, status=400)
                
            documents_content = "\n\n".join(documents_content)
                
        except Exception as e:
            print(f"Error processing documents: {str(e)}")
            return JsonResponse({'error': f'Error processing documents: {str(e)}'}, status=400)

        print(f"Generating blog content for topic: {topic}")
        print(f"Parameters - Tone: {tone}, Keywords: {user_keywords}")
        print(f"Successfully processed {successful_extractions} documents")
        print(f"Total content length: {len(documents_content)} characters")

        # Generate blog content using OpenAI
        try:
            result = generate_blog_content(
                topic=topic,
                tone=tone,
                target_keywords=user_keywords,
                documents_content=documents_content
            )
        except Exception as e:
            print(f"Error generating blog content: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)

        # Create Blog instance
        try:
            # First create the blog with required fields
            blog = Blog.objects.create(
                user=request.user,
                topic=topic,
                tone=tone,
                target_keywords=user_keywords,
                additional_keywords=result.get('keywords', {}).get('additional_keywords', [])
            )
            
            # Then update with generated content
            blog.blog_title = result.get('blog_title', topic)
            blog.blog_outline = result.get('blog_outline', '')
            blog.blog_draft = result.get('blog_draft', '')
            blog.save()
            
            # Associate selected documents with the blog
            blog.reference_documents.set(documents)
            
            print(f"Successfully created blog with ID: {blog.id}")

            # Return the created blog data
            return JsonResponse({
                'id': blog.id,
                'topic': blog.topic,
                'blog_title': blog.blog_title,
                'blog_outline': blog.blog_outline,
                'blog_draft': blog.blog_draft,
                'keywords': {
                    'user_keywords': user_keywords,
                    'additional_keywords': result.get('keywords', {}).get('additional_keywords', [])
                }
            }, status=201)

        except Exception as e:
            print(f"Error saving blog to database: {str(e)}")
            return JsonResponse({'error': f'Error saving blog: {str(e)}'}, status=500)

    except Exception as e:
        print(f"Unexpected error in create_blog view: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@csrf_exempt
def generate_blog(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            blog_id = data.get('blog_id')
            
            if not blog_id:
                return JsonResponse({'error': 'Blog ID is required'}, status=400)

            blog = Blog.objects.get(id=blog_id, user=request.user)
            
            # Get the selected documents
            documents = Document.objects.filter(id__in=data.get('document_ids', []), user=request.user)
            if not documents.exists():
                return JsonResponse({'error': 'No valid documents selected'}, status=400)

            # Extract content from all documents
            documents_content = ""
            for document in documents:
                try:
                    content = extract_text_from_file(document.file)
                    documents_content += f"\n\nDocument: {document.title}\n{content}"
                except Exception as e:
                    return JsonResponse({'error': f'Error processing document {document.title}: {str(e)}'}, status=500)

            # Generate blog content
            try:
                result = generate_blog_content(
                    topic=blog.topic,
                    tone=blog.tone,
                    target_keywords=blog.target_keywords,
                    documents_content=documents_content
                )

                # Update blog with generated content
                blog.blog_title = result['blog_title']
                blog.user_keywords = result['keywords']['user_keywords']
                blog.additional_keywords = result['keywords']['additional_keywords']
                blog.blog_outline = result['blog_outline']
                blog.blog_draft = result['blog_draft']
                blog.save()

                return JsonResponse({
                    'success': True,
                    'message': 'Blog generated successfully',
                    'blog': {
                        'id': blog.id,
                        'title': blog.blog_title,
                        'outline': blog.blog_outline
                    }
                })

            except Exception as e:
                return JsonResponse({'error': f'Error generating blog content: {str(e)}'}, status=500)

        except Blog.DoesNotExist:
            return JsonResponse({'error': 'Blog not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

@login_required
def edit_blog(request, blog_id):
    """
    Handle blog editing with form validation and content updates.
    """
    blog = get_object_or_404(Blog, id=blog_id, user=request.user)
    
    # Enhanced debug logging
    print("\n" + "="*80)
    print("BLOG DATA DEBUG LOG")
    print("="*80)
    print(f"Blog ID: {blog.id}")
    print(f"Title: {blog.blog_title}")
    print(f"Topic: {blog.topic}")
    print(f"Target Keywords: {blog.target_keywords}")
    print(f"Additional Keywords: {blog.additional_keywords}")
    print("\nOutline Data:")
    print(f"Type: {type(blog.blog_outline)}")
    print(f"Raw Value: {blog.blog_outline}")
    print("\nDraft Data:")
    print(f"Type: {type(blog.blog_draft)}")
    print(f"Raw Value: {blog.blog_draft}")
    print("="*80)
    
    if request.method == 'POST':
        form = BlogCreationForm(request.POST, instance=blog)
        if form.is_valid():
            blog = form.save(commit=False)
            blog.save()
            return JsonResponse({'success': True})
        else:
            return JsonResponse({'success': False, 'error': form.errors})

    # Create a context dictionary with the blog data
    context = {
        'blog': blog,
        'form': BlogCreationForm(instance=blog),
    }
    
    return render(request, 'core/edit_blog.html', context)

def format_blog_content(content):
    """Helper function to format blog content into HTML"""
    if not content:
        return ""

    if isinstance(content, str):
        try:
            content = json.loads(content)
        except json.JSONDecodeError:
            return f"<div class='content-section'>{content}</div>"

    html = ['<div class="content-container">']
    
    # Add Introduction
    if 'introduction' in content:
        html.append('<div class="section introduction-section">')
        html.append('<h2 class="section-title">Introduction</h2>')
        if isinstance(content['introduction'], dict):
            html.append(f'<div class="section-content">{content["introduction"]["content"]}</div>')
        else:
            html.append(f'<div class="section-content">{content["introduction"]}</div>')
        html.append('</div>')

    # Process main sections (H2 and H3)
    sections = [(k, v) for k, v in content.items() if k.startswith('H2_')]
    sections.sort(key=lambda x: int(x[0].split('_')[1]))

    for h2_key, h2_value in sections:
        html.append('<div class="section main-section">')
        if isinstance(h2_value, dict):
            html.append(f'<h2 class="section-title">{h2_value["content"]}</h2>')
        else:
            html.append(f'<h2 class="section-title">{h2_value}</h2>')
        
        # Find and add corresponding H3 sections
        h3_sections = [(k, v) for k, v in content.items() if k.startswith(f"{h2_key}_")]
        h3_sections.sort(key=lambda x: int(x[0].split('_')[-1]))
        
        for h3_key, h3_value in h3_sections:
            html.append('<div class="subsection">')
            if isinstance(h3_value, dict):
                html.append(f'<h3 class="subsection-title">{h3_value["content"]}</h3>')
                if "content" in h3_value:
                    html.append(f'<div class="subsection-content">{h3_value["content"]}</div>')
            else:
                html.append(f'<h3 class="subsection-title">{h3_value}</h3>')
            html.append('</div>')
        
        html.append('</div>')

    # Add Conclusion
    if 'conclusion' in content:
        html.append('<div class="section conclusion-section">')
        html.append('<h2 class="section-title">Conclusion</h2>')
        if isinstance(content['conclusion'], dict):
            html.append(f'<div class="section-content">{content["conclusion"]["content"]}</div>')
        else:
            html.append(f'<div class="section-content">{content["conclusion"]}</div>')
        html.append('</div>')

    html.append('</div>')
    return '\n'.join(html)

@login_required
@require_POST
def update_blog(request, blog_id):
    try:
        blog = Blog.objects.get(id=blog_id, user=request.user)
        
        # Get data from JSON body
        data = json.loads(request.body)
        
        # Update blog fields
        blog.blog_title = data.get('blog_title', blog.blog_title)
        blog.blog_outline = data.get('blog_outline', blog.blog_outline)
        blog.blog_draft = data.get('blog_draft', blog.blog_draft)
        blog.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Blog updated successfully'
        })
    except Blog.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Blog not found'
        }, status=404)
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@login_required
def download_blog(request, blog_id):
    try:
        blog = Blog.objects.get(id=blog_id, user=request.user)
        
        # Create a new Word document
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(blog.blog_title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add keywords section
        doc.add_heading('Keywords', level=2)
        keywords_para = doc.add_paragraph()
        keywords_para.add_run('User Keywords: ').bold = True
        keywords_para.add_run(', '.join(blog.user_keywords))
        keywords_para = doc.add_paragraph()
        keywords_para.add_run('Additional Keywords: ').bold = True
        keywords_para.add_run(', '.join(blog.additional_keywords))
        
        # Add outline section
        doc.add_heading('Outline', level=2)
        outline_para = doc.add_paragraph(blog.blog_outline)
        
        # Add draft section
        doc.add_heading('Draft', level=2)
        draft_para = doc.add_paragraph(blog.blog_draft)
        
        # Create a BytesIO object to save the document
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create the response
        response = HttpResponse(
            doc_io.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{blog.blog_title}.docx"'
        
        return response
        
    except Blog.DoesNotExist:
        messages.error(request, 'Blog not found')
        return redirect('dashboard')
    except Exception as e:
        messages.error(request, f'Error downloading blog: {str(e)}')
        return redirect('dashboard')

@login_required
@require_POST
def check_grammar_view(request):
    try:
        data = json.loads(request.body)
        text = data.get('text', '')
        
        if not text:
            return JsonResponse({'error': 'No text provided'}, status=400)
        
        suggestions = check_grammar(text)
        return JsonResponse({'suggestions': suggestions})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def manage_documents(request):
    """
    View function to display and manage user's uploaded documents.
    """
    documents = Document.objects.filter(user=request.user).order_by('-uploaded_at')
    return render(request, 'core/manage_documents.html', {'documents': documents})

@login_required
def delete_document(request, document_id):
    """
    View function to handle document deletion.
    """
    document = get_object_or_404(Document, id=document_id, user=request.user)
    try:
        document.delete()
        messages.success(request, 'Document deleted successfully!')
    except Exception as e:
        messages.error(request, f'Error deleting document: {str(e)}')
    return redirect('core:dashboard')
