import os
import openai
from django.conf import settings
import PyPDF2
from docx import Document as DocxDocument
import json
from io import BytesIO
import language_tool_python
import textwrap

def summarize_text(text, max_chars=10000):
    """Summarize text to fit within token limits."""
    if len(text) <= max_chars:
        return text
    
    # Split into paragraphs and select key ones
    paragraphs = text.split('\n\n')
    
    # Keep first and last paragraphs, and sample from middle
    summary_parts = []
    if paragraphs:
        # Always include first paragraph
        summary_parts.append(paragraphs[0])
        
        # Sample from middle paragraphs if there are more than 3
        if len(paragraphs) > 3:
            middle_count = min(5, len(paragraphs) - 2)  # Take up to 5 middle paragraphs
            step = len(paragraphs) // (middle_count + 1)
            for i in range(1, middle_count + 1):
                idx = i * step
                if idx < len(paragraphs):
                    summary_parts.append(paragraphs[idx])
        
        # Always include last paragraph if different from first
        if len(paragraphs) > 1 and paragraphs[-1] != paragraphs[0]:
            summary_parts.append(paragraphs[-1])
    
    summary = '\n\n'.join(summary_parts)
    
    # If still too long, truncate with ellipsis
    if len(summary) > max_chars:
        summary = summary[:max_chars-3] + '...'
    
    return summary

def extract_text_from_file(file):
    """Extract text content from different file types."""
    try:
        file_extension = os.path.splitext(file.name)[1].lower()
        
        if file_extension == '.pdf':
            return extract_text_from_pdf(file)
        elif file_extension in ['.doc', '.docx']:
            return extract_text_from_docx(file)
        else:
            # For txt files or other text-based formats
            content = ""
            try:
                file.seek(0)  # Reset file pointer
                content = file.read().decode('utf-8')
            except UnicodeDecodeError:
                try:
                    file.seek(0)
                    content = file.read().decode('latin-1')
                except:
                    file.seek(0)
                    content = file.read().decode('cp1252', errors='ignore')
            return content or "No content could be extracted"

    except Exception as e:
        print(f"Error in extract_text_from_file: {str(e)}")
        return "Error extracting text from file"

def extract_text_from_pdf(file):
    """Extract text from PDF files."""
    try:
        file.seek(0)  # Reset file pointer
        pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
        text = []
        
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            except Exception as e:
                print(f"Error extracting text from PDF page: {str(e)}")
                continue
        
        return "\n\n".join(text) if text else "No text could be extracted from PDF"
    
    except Exception as e:
        print(f"Error in extract_text_from_pdf: {str(e)}")
        return "Error extracting text from PDF"

def extract_text_from_docx(file):
    """Extract text from DOCX files."""
    try:
        file.seek(0)  # Reset file pointer
        doc = DocxDocument(BytesIO(file.read()))
        paragraphs = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return "\n\n".join(paragraphs) if paragraphs else "No text could be extracted from DOCX"
    
    except Exception as e:
        print(f"Error in extract_text_from_docx: {str(e)}")
        return "Error extracting text from DOCX"

def generate_blog_content(topic, tone, target_keywords, documents_content):
    """Generate blog content using OpenAI API."""
    # Initialize OpenAI client
    client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)

    # Convert target_keywords to list if it's a string
    if isinstance(target_keywords, str):
        target_keywords = [kw.strip() for kw in target_keywords.split(',') if kw.strip()]

    # Summarize documents_content to fit within token limits
    summarized_content = summarize_text(documents_content)
    print(f"Original content length: {len(documents_content)} chars")
    print(f"Summarized content length: {len(summarized_content)} chars")

    # Construct the prompt
    prompt = f"""You are a professional content writer and SEO expert. Create a high-quality, 100% human-written blog article that will outrank competing content. 

Topic: {topic}
Tone: {tone}
Target Keywords: {', '.join(target_keywords)}
Reference Documents Summary: {summarized_content}

Important Requirements:
1. Content must be 100% human-like - it should pass any AI detection tool
2. Research and analyze top-ranking blogs for this topic
3. Determine optimal word count based on competitor analysis
4. Include LSI keywords and semantic variations
5. Use natural language patterns and varied sentence structures
6. Include industry-specific terminology from the provided documents
7. Create an engaging, conversational flow
8. Implement advanced SEO strategies:
   - Strategic keyword placement
   - Proper heading hierarchy
   - Rich, engaging meta description
   - Internal linking suggestions
   - Featured snippet optimization

Structure the output as a JSON with:
{{
    "blog_title": "SEO-optimized, engaging title",
    "meta_description": "Compelling meta description for SEO",
    "keywords": {{
        "user_keywords": {json.dumps(target_keywords)},
        "additional_keywords": [],
        "lsi_keywords": [],
        "semantic_variations": []
    }},
    "word_count": "Recommended word count based on competitor analysis",
    "blog_outline": "Detailed outline with H2/H3 tags",
    "blog_draft": "Complete, human-like blog content",
    "seo_recommendations": {{
        "internal_linking": [],
        "featured_snippet_opportunities": [],
        "content_gaps": []
    }}
}}"""

    try:
        print(f"Sending request to OpenAI API with topic: {topic}")
        
        # Call OpenAI API with the new client format
        response = client.chat.completions.create(
            model="gpt-4-turbo-preview",
            messages=[
                {"role": "system", "content": "You are a professional content writer and SEO expert. Write in a completely human-like style that will pass any AI detection tool. Use natural language patterns, varied sentence structures, and industry-specific terminology."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.8,  # Increased for more human-like variation
            max_tokens=4000,
            response_format={"type": "json_object"}
        )

        print("Received response from OpenAI API")
        
        # Extract and parse the response
        content = response.choices[0].message.content
        print(f"Raw API response content: {content[:200]}...")
        
        try:
            # Parse the JSON response
            result = json.loads(content)
            
            # Validate the response structure
            required_keys = ['blog_title', 'meta_description', 'keywords', 'word_count', 'blog_outline', 'blog_draft', 'seo_recommendations']
            missing_keys = [key for key in required_keys if key not in result]
            if missing_keys:
                raise ValueError(f"Missing required keys in response: {missing_keys}")
            
            # Ensure keywords structure is correct
            if 'keywords' not in result or not isinstance(result['keywords'], dict):
                result['keywords'] = {
                    'user_keywords': target_keywords,
                    'additional_keywords': [],
                    'lsi_keywords': [],
                    'semantic_variations': []
                }
            
            print("Successfully validated response structure")
            return result
            
        except json.JSONDecodeError as e:
            print(f"JSON Decode Error: {str(e)}")
            # Create a fallback response
            return {
                'blog_title': topic,
                'meta_description': '',
                'keywords': {
                    'user_keywords': target_keywords,
                    'additional_keywords': [],
                    'lsi_keywords': [],
                    'semantic_variations': []
                },
                'word_count': '1500-2000 words',
                'blog_outline': 'Error generating outline',
                'blog_draft': 'Error generating draft',
                'seo_recommendations': {
                    'internal_linking': [],
                    'featured_snippet_opportunities': [],
                    'content_gaps': []
                }
            }

    except Exception as e:
        print(f"Error in generate_blog_content: {str(e)}")
        if hasattr(e, 'response'):
            print(f"OpenAI API Response: {e.response}")
        raise Exception(f"Error generating blog content: {str(e)}")

def check_grammar(text):
    """
    Check grammar in the given text and return suggestions.
    Returns a list of dictionaries containing grammar suggestions.
    """
    tool = language_tool_python.LanguageTool('en-US')
    matches = tool.check(text)
    
    suggestions = []
    for match in matches:
        suggestion = {
            'from': match.offset,
            'to': match.offset + match.errorLength,
            'message': match.message,
            'replacements': match.replacements,
            'context': match.context,
            'context_offset': match.contextOffset
        }
        suggestions.append(suggestion)
    
    return suggestions 